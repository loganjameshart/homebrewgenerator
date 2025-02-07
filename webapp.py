import os
import random
import pandas as pd
from bottle import Bottle, run, request, template, static_file, response, TEMPLATE_PATH
from docx import Document


# Make current working directory the webapp's directory, add it to Bottle.py path to be safe
dir_path = os.path.dirname(os.path.realpath(__file__))
os.chdir(dir_path)
TEMPLATE_PATH.insert(0, rf'{dir_path}\templates')


# Directory to store uploaded files temporarily
UPLOAD_DIR = 'uploads'
if not os.path.exists(UPLOAD_DIR):
    os.makedirs(UPLOAD_DIR)


# Generate Word doc from CSV file
def generate_document(csv_file_path):
    try:
        # Read the CSV file into a DataFrame
        df = pd.read_csv(csv_file_path, encoding="utf8")
        df.fillna('0', inplace=True)
        
        # Create a Word document
        export_document = Document()

        for column in df.columns:
            choice_list = df[column].tolist()
            while True:
                choice = random.choice(choice_list)
                if choice == '0':
                    continue
                else:
                    break

            export_document.add_heading(column)
            export_document.add_paragraph(choice)

        # Save the Word document as a downloadable file
        docx_filename = 'export.docx'
        docx_path = os.path.join(UPLOAD_DIR, docx_filename)
        export_document.save(docx_path)
        
        # Return the document as a downloadable file
        return static_file(docx_filename, root=UPLOAD_DIR, download=True)

    except Exception as e:
        return f"Error generating document: {e}"


'''~~~Bottle Functions~~~'''
# Create a Bottle web app
app = Bottle()


# Define route to show the upload form
@app.route('/')
def upload_form():
    return template('upload_form.tpl')


# Route to handle the file upload and processing
@app.route('/upload', method='POST')
def do_upload():
    upload = request.files.get('csv_file')
    if not upload:
        return "No file uploaded"
    
    # Save the uploaded file to the uploads folder
    file_path = os.path.join(UPLOAD_DIR, upload.filename)
    upload.save(file_path)
    
    # Process the file and generate the document
    return generate_document(file_path)


# Route to serve static files (if needed, for images, CSS, etc.)
@app.route('/static/<filename>')
def send_static(filename):
    return static_file(filename, root='static')


if __name__ == '__main__':
    # Main Bottle function
    run(app, host='localhost', port=8080, reloader=True)
